import docx
import re
import io

def extract_table_from_upload(file):
    doc = docx.Document(io.BytesIO(file.read()))
    res = {
        "审理法院": "未提取",
        "案号": "未提取",
        "案由": "未提取",
        "原告": "未提取",
        "被告": "未提取",
        "全文": ""
    }

    # 遍历所有表格，适配单元格带空格：案  号、审 理 法 院
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            row_key = "".join(cells).replace(" ", "")

            if "审理法院" in row_key:
                res["审理法院"] = cells[-1].strip()
            elif "案号" in row_key:
                res["案号"] = cells[-1].strip()
            elif "案由" in row_key:
                val = cells[-1].strip()
                res["案由"] = val.split("/")[-1]
            elif "原告" in row_key and "被告" not in row_key:
                res["原告"] = cells[-1].strip()
            elif "被告" in row_key and "原告" not in row_key:
                res["被告"] = cells[-1].strip()

    # 提取全文段落
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    res["全文"] = "\n".join(full_text)

    # 正文兜底补案由
    if res["案由"] == "未提取":
        pat = r"原告.*?诉.*?被告.*?([\u4e00-\u9fa5]+纠纷)"
        m = re.search(pat, res["全文"])
        if m:
            res["案由"] = m.group(1)

    return res
